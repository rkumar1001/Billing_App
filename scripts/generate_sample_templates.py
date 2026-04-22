"""Generate a realistic Bella Vista style Excel + Word pair for manual testing.

Run:
    cd /home/rajat/Desktop/n8n\\ workflows/Billing_App
    .venv/bin/python scripts/generate_sample_templates.py

Outputs land in  sample_templates/Bella Vista/March/
so the user can point the Profile Wizard at them.
"""
from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "sample_templates" / "Bella Vista" / "March"
OUT.mkdir(parents=True, exist_ok=True)


# ---------- EXCEL ---------------------------------------------------------
def build_excel() -> Path:
    path = OUT / "Breakdown of Services - Bella Vista - March.xlsx"
    wb = Workbook()
    ws = wb.active
    ws.title = "Breakdown"

    bold = Font(bold=True)
    yellow = PatternFill("solid", start_color="FFF2CC", end_color="FFF2CC")
    center = Alignment(horizontal="center", vertical="center")
    left = Alignment(horizontal="left", vertical="center")
    right = Alignment(horizontal="right", vertical="center")

    # Row 1: headings and invoice month number cell
    ws["A1"] = "Toll Brothers"
    ws["A1"].font = bold
    ws["D1"] = "BREAKDOWN OF SERVICES"
    ws["D1"].font = Font(bold=True, size=14)
    ws.merge_cells("D1:G1")
    ws["D1"].alignment = center
    ws["J1"] = "International Security Services"
    ws["J1"].font = bold
    ws["J1"].alignment = right
    # Month number (top-right, stands alone) — this is what the wizard will
    # map to `month_number`.
    ws["K2"] = 3
    ws["K2"].font = Font(bold=True, size=14)
    ws["K2"].alignment = center
    # Invoice date (yellow highlighted per PDF) — mapped to `invoice_date_excel`.
    ws["K3"] = date(2026, 3, 6)
    ws["K3"].fill = yellow
    ws["K3"].alignment = center
    ws["K3"].number_format = "dddd, mmmm d, yyyy"

    ws["A3"] = "BELLA VISTA COMMUNITY"
    ws["A3"].font = bold

    # Alphabet row (row 4) just for reference, matching the PDF.
    for i, letter in enumerate("ABCDEFGHIJ", start=1):
        ws.cell(row=4, column=i, value=letter).alignment = center

    # Column headers at row 5.
    headers = [
        "ITEM", "Date of Work", "Hours per day", "Number of Guards",
        "Regular Number of Hours", "Hourly Rate", "Daily Amount",
        "Additional Hours", "Additional Amount", "Description of Work",
    ]
    for col, text in enumerate(headers, start=1):
        cell = ws.cell(row=5, column=col, value=text)
        cell.font = bold
        cell.alignment = center

    # Data rows (31 rows for March).
    hours_per_day = 11
    guards = 1
    rate = 20
    n_days = 31
    for i in range(1, n_days + 1):
        row = 5 + i
        ws.cell(row=row, column=1, value=i)
        d = ws.cell(row=row, column=2, value=date(2026, 3, i))
        d.number_format = "dddd, mmmm d, yyyy"
        ws.cell(row=row, column=3, value=hours_per_day)
        ws.cell(row=row, column=4, value=guards)
        # Regular hours = hours per day * guards
        ws.cell(row=row, column=5, value=f"=C{row}*D{row}")
        ws.cell(row=row, column=6, value=rate)
        amt = ws.cell(row=row, column=7, value=f"=E{row}*F{row}")
        amt.number_format = '"$"#,##0.00'
        if i == 1:
            ws.cell(row=row, column=10, value="6:30 AM to 5:30 PM")

    # Totals row (after the last day).
    totals_row = 5 + n_days + 1
    bd = ws.cell(row=totals_row, column=2, value="Breakdown")
    bd.font = bold
    ws.cell(row=totals_row, column=5, value=f"=SUM(E6:E{5 + n_days})")
    r_cell = ws.cell(row=totals_row, column=6, value=rate)
    r_cell.number_format = '"$"#,##0.00'
    g_cell = ws.cell(row=totals_row, column=7, value=f"=SUM(G6:G{5 + n_days})")
    g_cell.number_format = '"$"#,##0.00'
    ws.cell(row=totals_row, column=8, value=0)
    ws.cell(row=totals_row, column=9, value=0).number_format = '"$"#,##0.00'
    grand = ws.cell(row=totals_row, column=10, value=f"=G{totals_row}+I{totals_row}")
    grand.number_format = '"$"#,##0.00'

    # "TOTAL HOURS" footer row.
    foot_row = totals_row + 1
    ws.cell(row=foot_row, column=3, value="TOTAL HOURS").font = bold
    ws.cell(row=foot_row, column=5, value=f"=E{totals_row}")

    # Column widths for readability.
    widths = [6, 28, 12, 14, 16, 12, 14, 12, 14, 20, 12]
    for i, w in enumerate(widths, start=1):
        ws.column_dimensions[get_column_letter(i)].width = w

    wb.save(str(path))
    return path


# ---------- WORD ----------------------------------------------------------
def build_word() -> Path:
    path = OUT / "Invoice Bella16.docx"
    doc = Document()

    # Header meta table (left-side block in the PDF).
    meta = doc.add_table(rows=7, cols=2)
    meta.style = "Light Grid"
    meta_rows = [
        ("Invoice", "Bella16"),
        ("Invoice Date", "March 6, 2026"),
        ("Name", "Toll Brothers"),
        ("Job Site", "12075 Red Hawk"),
        ("Address", "Porter Ranch"),
        ("State", "CA"),
        ("Zip", "91326"),
    ]
    for i, (k, v) in enumerate(meta_rows):
        meta.rows[i].cells[0].text = k
        meta.rows[i].cells[1].text = v
        for p in meta.rows[i].cells[0].paragraphs:
            for r in p.runs:
                r.bold = True

    doc.add_paragraph("")
    title = doc.add_paragraph()
    run = title.add_run("INTERNATIONAL SECURITY SERVICES")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0xC9, 0x8A, 0x00)
    doc.add_paragraph("PPO License: 16863")
    doc.add_paragraph("")

    # Service-line table (No. / Description / Billing Period / Hourly Rate / Total Hours).
    service = doc.add_table(rows=6, cols=5)
    service.style = "Light Grid"
    hdr = service.rows[0].cells
    hdr[0].text = "No."
    hdr[1].text = "Description"
    hdr[2].text = "Billing Period"
    hdr[3].text = "Hourly Rate"
    hdr[4].text = "Total Hours"
    for c in hdr:
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True

    first_row = service.rows[1].cells
    first_row[0].text = "1."
    first_row[1].text = "Guard Service"
    first_row[2].text = "3/1/2026 – 3/31/2026"
    first_row[3].text = "$20"
    # Per PDF this cell shows "308" — stale value from Feb (28 × 11) in the
    # template before the user updates it. Leave as-is; the app overwrites it.
    first_row[4].text = "308"

    # Grand total table (right-aligned block with "Comment" on the left).
    doc.add_paragraph("")
    totals = doc.add_table(rows=2, cols=2)
    totals.style = "Light Grid"
    totals.rows[0].cells[0].text = "Comment"
    totals.rows[0].cells[1].text = ""
    totals.rows[1].cells[0].text = "Grand Total"
    totals.rows[1].cells[1].text = "$6,820"
    for p in totals.rows[1].cells[0].paragraphs:
        for r in p.runs:
            r.bold = True

    doc.add_paragraph("")
    doc.add_paragraph(
        "Please make check payable to International Security Services"
    )
    doc.add_paragraph(
        "And mail checks to 5877 Evening Sky Drive. Simi Valley, CA 93063"
    )

    doc.save(str(path))
    return path


def main() -> None:
    xl = build_excel()
    wd = build_word()
    print(f"Excel: {xl}")
    print(f"Word : {wd}")
    print()
    print("Profile Wizard mapping hints:")
    print("  Files:")
    print(f"    Excel template: {xl}")
    print(f"    Word template : {wd}")
    print("    Invoice prefix: Bella")
    print("    Next invoice number: 17")
    print("    Hourly rate: 20")
    print("    Hours per day: 11")
    print("    Guards: 1")
    print()
    print("  Excel mapping (sheet 'Breakdown'):")
    print("    Month Number          -> K2")
    print("    Invoice Date (Excel)  -> K3")
    print("    Dates Column Start    -> B6   (fills B6..B36 with April 1..30, clears 37)")
    print("    Total Hours (Excel)   -> E37  (optional - formula recalculates in Excel)")
    print("    Grand Total (Excel)   -> J37  (optional - formula recalculates in Excel)")
    print()
    print("  Word mapping (table / row / col / paragraph):")
    print("    Invoice Number          -> 0 / 0 / 1 / 0")
    print("    Invoice Date            -> 0 / 1 / 1 / 0")
    print("    Billing Period (one cell) -> 1 / 1 / 2 / 0")
    print("    Total Hours             -> 1 / 1 / 4 / 0")
    print("    Grand Total             -> 2 / 1 / 1 / 0")


if __name__ == "__main__":
    main()
