from __future__ import annotations

from pathlib import Path
from tkinter import ttk
from typing import Callable

import customtkinter as ctk
from docx import Document
from openpyxl import load_workbook
from openpyxl.utils import get_column_letter

from ..services.auto_detect import WordLoc


class ExcelCellPicker(ctk.CTkToplevel):
    """Modal that shows every sheet + cell of an Excel file and asks the user
    to click the cell they meant. Returns `(sheet, cell_address)`.
    """

    def __init__(self, parent, xlsx_path: Path, title: str = "Pick cell") -> None:
        super().__init__(parent)
        self.title(title)
        self.geometry("900x600")
        self.transient(parent)
        self.selected_cell: tuple[str, str] | None = None

        wb = load_workbook(filename=str(xlsx_path), data_only=True)
        try:
            self._sheets = list(wb.sheetnames)
            self._data: dict[str, list[list]] = {}
            for name in self._sheets:
                ws = wb[name]
                rows: list[list] = []
                max_rows = min(60, ws.max_row or 60)
                max_cols = min(16, ws.max_column or 16)
                for r in range(1, max_rows + 1):
                    row: list = []
                    for c in range(1, max_cols + 1):
                        row.append(ws.cell(row=r, column=c).value)
                    rows.append(row)
                self._data[name] = rows
        finally:
            wb.close()

        self.grid_rowconfigure(1, weight=1)
        self.grid_columnconfigure(0, weight=1)

        top = ctk.CTkFrame(self, fg_color="transparent")
        top.grid(row=0, column=0, sticky="ew", padx=12, pady=(12, 4))
        ctk.CTkLabel(
            top,
            text="Double-click the cell you want to map, or select a row and type a cell address.",
            font=ctk.CTkFont(size=12),
        ).pack(side="left")

        self._sheet_var = ctk.StringVar(value=self._sheets[0])
        ctk.CTkOptionMenu(
            top, variable=self._sheet_var, values=self._sheets,
            command=lambda *_: self._populate(),
        ).pack(side="right")

        self._tree = ttk.Treeview(self, show="headings", height=25)
        self._tree.grid(row=1, column=0, sticky="nsew", padx=12, pady=4)
        self._tree.bind("<Double-1>", self._on_double_click)

        bottom = ctk.CTkFrame(self, fg_color="transparent")
        bottom.grid(row=2, column=0, sticky="ew", padx=12, pady=(4, 12))
        ctk.CTkLabel(bottom, text="Cell:").pack(side="left")
        self._cell_var = ctk.StringVar()
        ctk.CTkEntry(bottom, textvariable=self._cell_var, width=120).pack(side="left", padx=6)
        ctk.CTkButton(bottom, text="Cancel", command=self.destroy,
                      fg_color="transparent", border_width=1).pack(side="right", padx=4)
        ctk.CTkButton(bottom, text="OK", command=self._confirm).pack(side="right", padx=4)

        self._populate()

    def _populate(self) -> None:
        for iid in self._tree.get_children():
            self._tree.delete(iid)
        name = self._sheet_var.get()
        rows = self._data.get(name, [])
        if not rows:
            return
        max_cols = max(len(r) for r in rows)
        columns = ["row"] + [get_column_letter(c) for c in range(1, max_cols + 1)]
        self._tree.configure(columns=columns)
        self._tree.heading("row", text="#")
        self._tree.column("row", width=36, anchor="center", stretch=False)
        for letter in columns[1:]:
            self._tree.heading(letter, text=letter)
            self._tree.column(letter, width=100, anchor="w")
        for i, row in enumerate(rows, start=1):
            display = [i] + [("" if v is None else str(v))[:24] for v in row]
            self._tree.insert("", "end", values=display)

    def _on_double_click(self, event) -> None:
        region = self._tree.identify("region", event.x, event.y)
        if region != "cell":
            return
        column_id = self._tree.identify_column(event.x)
        row_iid = self._tree.identify_row(event.y)
        if not row_iid or not column_id:
            return
        col_idx = int(column_id.replace("#", "")) - 1
        if col_idx <= 0:
            return
        row_number = self._tree.index(row_iid) + 1
        column_letter = get_column_letter(col_idx)
        self._cell_var.set(f"{column_letter}{row_number}")

    def _confirm(self) -> None:
        cell = self._cell_var.get().strip().upper()
        if not cell:
            self.destroy()
            return
        self.selected_cell = (self._sheet_var.get(), cell)
        self.destroy()


class WordLocationPicker(ctk.CTkToplevel):
    """Modal that shows the Word document's table + paragraph tree and asks
    the user to select one. Returns a `WordLoc`.
    """

    def __init__(self, parent, docx_path: Path, title: str = "Pick Word location") -> None:
        super().__init__(parent)
        self.title(title)
        self.geometry("900x600")
        self.transient(parent)
        self.selected_location: WordLoc | None = None

        doc = Document(str(docx_path))
        self.grid_rowconfigure(1, weight=1)
        self.grid_columnconfigure(0, weight=1)

        ctk.CTkLabel(
            self,
            text="Double-click the cell that contains the value you want to update.",
            font=ctk.CTkFont(size=12),
        ).grid(row=0, column=0, sticky="w", padx=12, pady=(12, 2))

        self._tree = ttk.Treeview(self, columns=("text",), show="tree headings", height=25)
        self._tree.heading("#0", text="Location")
        self._tree.heading("text", text="Text")
        self._tree.column("#0", width=260)
        self._tree.column("text", width=560)
        self._tree.grid(row=1, column=0, sticky="nsew", padx=12, pady=4)
        self._tree.bind("<Double-1>", self._on_double_click)

        self._locations: dict[str, WordLoc] = {}

        for t_idx, table in enumerate(doc.tables):
            t_node = self._tree.insert(
                "", "end", text=f"Table {t_idx}", open=True, values=("",),
            )
            for r_idx, row in enumerate(table.rows):
                for c_idx, cell in enumerate(row.cells):
                    for p_idx, paragraph in enumerate(cell.paragraphs):
                        label = f"t{t_idx} r{r_idx} c{c_idx} p{p_idx}"
                        iid = self._tree.insert(
                            t_node, "end",
                            text=label,
                            values=(paragraph.text[:80],),
                        )
                        self._locations[iid] = WordLoc(t_idx, r_idx, c_idx, p_idx)

        bottom = ctk.CTkFrame(self, fg_color="transparent")
        bottom.grid(row=2, column=0, sticky="ew", padx=12, pady=(4, 12))
        ctk.CTkButton(bottom, text="Cancel", command=self.destroy,
                      fg_color="transparent", border_width=1).pack(side="right", padx=4)
        ctk.CTkButton(bottom, text="OK", command=self._confirm).pack(side="right", padx=4)

    def _on_double_click(self, event) -> None:
        iid = self._tree.identify_row(event.y)
        if iid in self._locations:
            self.selected_location = self._locations[iid]
            self.destroy()

    def _confirm(self) -> None:
        sel = self._tree.selection()
        if sel and sel[0] in self._locations:
            self.selected_location = self._locations[sel[0]]
        self.destroy()
