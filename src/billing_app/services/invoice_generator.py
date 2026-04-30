"""Folder-in, folder-out invoice duplicator.

Given a source folder (e.g. `.../Bella Vista/March/`) and a target month,
this module:
  1. Copies the whole folder next to itself with the target month name.
  2. Renames the Excel and Word files to reflect the new month and the next
     invoice number (Bella16 → Bella17).
  3. Auto-detects field locations on first contact (cached in
     `.billingapp.json` inside the source folder) and applies updates in the
     copy: month number, invoice date, dates column in Excel; invoice number,
     invoice date, billing period, total hours, grand total in Word.

User-provided values (target month, invoice date, and optional overrides for
rate / hours / guards) drive the updates; everything else is inferred.
"""
from __future__ import annotations

import platform
import re
import shutil
from dataclasses import dataclass, field
from datetime import date
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml.ns import qn
from openpyxl import load_workbook
from openpyxl.utils.cell import coordinate_from_string

from .auto_detect import (
    DetectionCache,
    ExcelDetection,
    WordDetection,
    WordLoc,
    detect_excel,
    detect_word,
    load_cache,
    save_cache,
)
from .auxiliary_excels import update_g703, update_payment_request_form
from .calendar_util import MONTH_NAMES, days_in_month, month_name
from .folder_analyzer import (
    FolderAnalysis, analyze, rename_for_target, swap_any_month_token,
)
from .xls_to_xlsx import XlsConvertError, convert_xls_to_xlsx


@dataclass
class GenerateRequest:
    source_folder: Path
    target_month: int
    target_year: int
    invoice_date: date
    target_folder_name: str | None = None
    hourly_rate: float | None = None
    hours_per_day: float | None = None
    guards: int | None = None
    explicit_invoice_number: int | None = None
    excel_overrides: ExcelDetection | None = None
    word_overrides: WordDetection | None = None
    overwrite: bool = False


@dataclass
class GenerateResult:
    copied_folder: Path
    excel_path: Path | None
    word_path: Path | None
    invoice_number: str
    invoice_number_int: int
    total_hours: float
    grand_total: float
    analysis: FolderAnalysis
    excel_detection: ExcelDetection
    word_detection: WordDetection
    unresolved_excel: list[str] = field(default_factory=list)
    unresolved_word: list[str] = field(default_factory=list)
    auxiliary_files: list[tuple[Path, str]] = field(default_factory=list)


class GeneratorError(Exception):
    pass


def preview(source: Path) -> tuple[FolderAnalysis, ExcelDetection | None, WordDetection | None]:
    """Run analysis + detection without making any changes. Used for UI."""
    analysis = analyze(Path(source))

    excel_det: ExcelDetection | None = None
    word_det: WordDetection | None = None

    cache = load_cache(analysis.folder)
    if cache:
        excel_det = cache.excel
        word_det = cache.word

    if analysis.excel_path:
        try:
            detected = detect_excel(analysis.excel_path)
        except Exception as e:  # noqa: BLE001
            detected = ExcelDetection()
            detected.warnings.append(f"Excel detection failed: {e}")
        if excel_det is None:
            excel_det = detected
        else:
            for attr in (
                "sheet", "month_number_cell", "invoice_date_cell",
                "dates_column_start", "dates_row_end",
            ):
                if getattr(excel_det, attr) is None:
                    setattr(excel_det, attr, getattr(detected, attr))
            # Always refresh derived fields from the live detection.
            excel_det.detected_month = detected.detected_month
            excel_det.detected_year = detected.detected_year
            excel_det.hourly_rate = detected.hourly_rate
            excel_det.hours_per_day = detected.hours_per_day
            excel_det.guards = detected.guards

    if analysis.word_path and analysis.invoice_number is not None:
        invoice_str = f"{analysis.invoice_prefix}{analysis.invoice_number}"
        try:
            detected_word = detect_word(analysis.word_path, invoice_str)
        except Exception as e:  # noqa: BLE001
            detected_word = WordDetection()
            detected_word.warnings.append(f"Word detection failed: {e}")
        if word_det is None:
            word_det = detected_word
        else:
            for attr in ("invoice_number", "invoice_date", "billing_period",
                          "total_hours", "grand_total"):
                if getattr(word_det, attr) is None:
                    setattr(word_det, attr, getattr(detected_word, attr))

    return analysis, excel_det, word_det


def generate(req: GenerateRequest) -> GenerateResult:
    source = Path(req.source_folder)
    if not source.is_dir():
        raise GeneratorError(f"Source folder does not exist: {source}")

    analysis = analyze(source)
    missing = analysis.missing()
    if missing:
        raise GeneratorError(
            "Source folder is missing required items: " + ", ".join(missing)
        )

    # Decide target folder name.
    target_name = req.target_folder_name
    if not target_name:
        target_name = _swap_month_in_name(
            source.name, analysis.folder_month_token, req.target_month
        )
    target_folder = source.parent / target_name
    if target_folder.exists():
        if not req.overwrite:
            raise GeneratorError(
                f"Target folder already exists: {target_folder}. "
                f"Enable overwrite to replace."
            )
        shutil.rmtree(target_folder)

    shutil.copytree(source, target_folder)

    # Drop any stale PDF exports copied over from the source folder.
    for item in target_folder.iterdir():
        if item.is_file() and item.suffix.lower() == ".pdf":
            try:
                item.unlink()
            except OSError:
                pass

    # Determine next invoice number.
    invoice_int = (
        req.explicit_invoice_number
        if req.explicit_invoice_number is not None
        else analysis.next_invoice_number
    )

    # Upgrade legacy .xls -> .xlsx INSIDE the target folder using a
    # real-fidelity backend (MS Excel COM on Windows, LibreOffice elsewhere).
    # The original .xls in the source stays untouched; the target's .xls is
    # replaced by an equivalent .xlsx with all formatting/formulas intact.
    if analysis.excel_is_legacy_xls and analysis.excel_path is not None:
        legacy_in_target = target_folder / analysis.excel_path.name
        if legacy_in_target.exists():
            new_path_in_target = legacy_in_target.with_suffix(".xlsx")
            try:
                convert_xls_to_xlsx(legacy_in_target, new_path_in_target)
            except XlsConvertError as e:
                # Roll back the partially-built target folder so the user's
                # tree is left clean, then surface a clear error.
                try:
                    shutil.rmtree(target_folder)
                except OSError:
                    pass
                raise GeneratorError(
                    f"Could not convert legacy .xls to .xlsx automatically: "
                    f"{e}"
                ) from e
            try:
                legacy_in_target.unlink()
            except OSError:
                pass
            analysis.excel_path = new_path_in_target

    new_excel, new_word = rename_for_target(
        analysis, target_folder, req.target_month, req.target_year, invoice_int,
    )

    # Resolve detection: overrides > cache > fresh detection.
    # For detection we read whichever Excel/Word file is currently in the
    # target folder (handles the legacy-xls conversion and post-rename paths
    # transparently).
    cache = load_cache(source)
    excel_det = req.excel_overrides or (cache.excel if cache else None)
    excel_for_detect = new_excel or analysis.excel_path
    if excel_det is None and excel_for_detect is not None:
        excel_det = detect_excel(excel_for_detect)
    if excel_det is None:
        excel_det = ExcelDetection()

    word_det = req.word_overrides or (cache.word if cache else None)
    word_for_detect = new_word or analysis.word_path
    if word_det is None and word_for_detect is not None and analysis.invoice_number is not None:
        word_det = detect_word(
            word_for_detect,
            f"{analysis.invoice_prefix}{analysis.invoice_number}",
        )
    if word_det is None:
        word_det = WordDetection()

    # Compute totals.
    n_days = days_in_month(req.target_year, req.target_month)
    rate = _first_not_none(
        req.hourly_rate, excel_det.hourly_rate, 0.0,
    )
    hpd = _first_not_none(
        req.hours_per_day, excel_det.hours_per_day, 0.0,
    )
    guards = int(_first_not_none(req.guards, excel_det.guards, 1))

    total_hours = round(n_days * hpd * guards, 2)
    grand_total = round(total_hours * rate, 2)

    # Default form (e.g. "HART24"); we'll adapt to the doc's spacing
    # (e.g. "HART 24") below if the source used a space.
    invoice_number_str = f"{analysis.invoice_prefix}{invoice_int}"
    if word_det and word_det.invoice_number and " " in word_det.invoice_number.match_text:
        invoice_number_str = f"{analysis.invoice_prefix} {invoice_int}"

    # Apply to Excel copy.
    unresolved_excel: list[str] = []
    if new_excel is not None:
        unresolved_excel = _update_excel(
            new_excel,
            excel_det,
            source_month=analysis.source_month or excel_det.detected_month,
            target_year=req.target_year,
            target_month=req.target_month,
            invoice_date=req.invoice_date,
        )

    # Apply to Word copy.
    unresolved_word: list[str] = []
    if new_word is not None:
        period_string = _format_period(req.target_year, req.target_month)
        invoice_date_string = _format_invoice_date(req.invoice_date)
        from_string, to_string = _format_period_endpoints(
            req.target_year, req.target_month,
        )
        unresolved_word = _update_word(
            new_word,
            word_det,
            values={
                "invoice_number": invoice_number_str,
                "invoice_date": invoice_date_string,
                "billing_period": period_string,
                "from_date": from_string,
                "to_date": to_string,
                "total_hours": _format_number(total_hours),
                "grand_total": _format_currency(grand_total),
            },
        )

    # Process auxiliary spreadsheets (PRF, G703) found alongside the main
    # breakdown. We convert any .xls to .xlsx in place, swap the month
    # token in the filename, and apply the role-specific edit recipe.
    auxiliary_results: list[tuple[Path, str]] = []
    for aux_src_path, role, aux_legacy in analysis.auxiliary_excels:
        aux_in_target = target_folder / aux_src_path.name
        if not aux_in_target.exists():
            continue
        # Convert .xls -> .xlsx in target if needed.
        if aux_legacy:
            try:
                xlsx_path = aux_in_target.with_suffix(".xlsx")
                convert_xls_to_xlsx(aux_in_target, xlsx_path)
                try:
                    aux_in_target.unlink()
                except OSError:
                    pass
                aux_in_target = xlsx_path
            except XlsConvertError:
                continue  # skip this auxiliary on conversion failure
        # Rename: swap any month token found in the filename. Auxiliary
        # forms often have stale month tags ("...February.xlsx") that
        # don't match the folder's month token, so scan all months.
        new_name = swap_any_month_token(aux_in_target.name, req.target_month)
        renamed = aux_in_target
        if new_name != aux_in_target.name:
            renamed = aux_in_target.with_name(new_name)
            try:
                aux_in_target.rename(renamed)
            except OSError:
                renamed = aux_in_target
        # Apply role-specific edits.
        try:
            if role == "payment_request_form":
                update_payment_request_form(
                    renamed,
                    target_year=req.target_year,
                    target_month=req.target_month,
                    invoice_date=req.invoice_date,
                )
            elif role == "g703":
                update_g703(
                    renamed,
                    target_year=req.target_year,
                    target_month=req.target_month,
                    invoice_date=req.invoice_date,
                    invoice_number_str=invoice_number_str,
                )
        except Exception:  # noqa: BLE001
            # Don't fail the whole generate if one auxiliary file blows up.
            pass
        auxiliary_results.append((renamed, role))

    # Write cache back into the SOURCE folder so next run is fast.
    try:
        if analysis.excel_path and analysis.word_path:
            save_cache(source, DetectionCache(excel=excel_det, word=word_det))
    except Exception:  # noqa: BLE001
        pass

    return GenerateResult(
        copied_folder=target_folder,
        excel_path=new_excel,
        word_path=new_word,
        invoice_number=invoice_number_str,
        invoice_number_int=invoice_int,
        total_hours=total_hours,
        grand_total=grand_total,
        analysis=analysis,
        excel_detection=excel_det,
        word_detection=word_det,
        unresolved_excel=unresolved_excel,
        unresolved_word=unresolved_word,
        auxiliary_files=auxiliary_results,
    )


# ---------------------------------------------------------- helpers --------

def _swap_month_in_name(name: str, token: str, target_month: int) -> str:
    import re as _re
    target_name = MONTH_NAMES[target_month - 1]
    if token:
        return _re.sub(_re.escape(token), target_name, name, flags=_re.IGNORECASE)
    return target_name  # folder didn't contain a month name; use raw target name


def _first_not_none(*values: Any) -> Any:
    for v in values:
        if v is not None and v != 0:
            return v
    for v in values:
        if v is not None:
            return v
    return None


def _update_excel(
    path: Path,
    det: ExcelDetection,
    *,
    source_month: int | None,
    target_year: int,
    target_month: int,
    invoice_date: date,
) -> list[str]:
    unresolved: list[str] = []
    wb = load_workbook(filename=str(path))
    try:
        ws = wb[det.sheet] if det.sheet and det.sheet in wb.sheetnames else wb.active

        if det.month_number_cell:
            ws[det.month_number_cell] = target_month
        else:
            unresolved.append("month_number_cell")

        if det.invoice_date_cell:
            ws[det.invoice_date_cell] = invoice_date
            ws[det.invoice_date_cell].number_format = "dddd, mmmm d, yyyy"
        else:
            unresolved.append("invoice_date_cell")

        if det.dates_column_start:
            col_letter, start_row = coordinate_from_string(det.dates_column_start)
            col_idx = _col_to_index(col_letter)
            target_days = days_in_month(target_year, target_month)
            if det.formula_dates:
                # Template uses =TEXT(DATE(YYYY, $K$_, A_)) formulas — the
                # date column auto-regenerates from the month cell we just
                # wrote above. We only need to (a) extend or trim data rows
                # so totals cover exactly target_days, and (b) swap the
                # hardcoded year if the user moved into a new year.
                _adjust_formula_rows(
                    ws,
                    start_row=start_row,
                    last_row=det.dates_row_end or (start_row + target_days - 1),
                    target_days=target_days,
                )
                if det.detected_year and det.detected_year != target_year:
                    _swap_year_in_formulas(
                        ws, det.detected_year, target_year,
                    )
            elif det.text_dates and det.text_date_format:
                # Template uses plain text strings for dates ("Wednesday,
                # April-1 ,2026" etc.). Re-render each row using the
                # detected strftime format.
                for i in range(target_days):
                    d = date(target_year, target_month, i + 1)
                    ws.cell(
                        row=start_row + i, column=col_idx,
                        value=d.strftime(det.text_date_format),
                    )
                # Trim leftover rows from a longer source month.
                source_days = det.dates_row_end and (det.dates_row_end - start_row + 1) or target_days
                if source_days > target_days:
                    for i in range(target_days, source_days):
                        ws.cell(row=start_row + i, column=col_idx).value = None
                # Extend other columns (hours/guards/rate/etc.) for newly
                # added rows by copying the last source row's values forward.
                if target_days > source_days:
                    last_src_row = start_row + source_days - 1
                    for i in range(source_days, target_days):
                        new_row = start_row + i
                        for col in range(1, ws.max_column + 1):
                            if col == col_idx:
                                continue
                            src = ws.cell(row=last_src_row, column=col)
                            dst = ws.cell(row=new_row, column=col)
                            if src.value is not None:
                                dst.value = src.value
                                _copy_cell_style(src, dst)
            else:
                # Write new-month dates literally.
                for i in range(target_days):
                    d = date(target_year, target_month, i + 1)
                    cell = ws.cell(row=start_row + i, column=col_idx, value=d)
                    cell.number_format = "dddd, mmmm d, yyyy"
                # Clear trailing rows when source month was longer.
                source_days_guess = det.dates_row_end and (det.dates_row_end - start_row + 1)
                source_days = source_days_guess or (
                    days_in_month(target_year, source_month) if source_month else 31
                )
                if source_days > target_days:
                    for i in range(target_days, source_days):
                        ws.cell(row=start_row + i, column=col_idx).value = None
        else:
            unresolved.append("dates_column_start")

        wb.save(str(path))
    finally:
        wb.close()
    return unresolved


def _update_word(
    path: Path,
    det: WordDetection,
    *,
    values: dict[str, str],
) -> list[str]:
    unresolved: list[str] = []
    doc = Document(str(path))

    # Gather leaf paragraphs (the same set the detector used). Skipping
    # container <w:p> wrappers means we don't accidentally also rewrite a
    # super-paragraph that encloses text frames.
    leaf_paragraphs = list(_leaf_paragraph_elements(doc))

    # Hartford-style invoices use a multi-row line-items table where every
    # row repeats the same from/to dates and has its own per-row total. The
    # user wants only the dates updated; per-row hours/amounts/totals stay.
    hartford_style = det.from_date is not None or det.to_date is not None

    fields_to_apply = ["invoice_number", "invoice_date"]
    if det.billing_period is not None:
        fields_to_apply.append("billing_period")
    if det.from_date is not None:
        fields_to_apply.append("from_date")
    if det.to_date is not None:
        fields_to_apply.append("to_date")
    # Only update totals when this is NOT a multi-row Hartford-style doc.
    if not hartford_style:
        fields_to_apply.extend(["total_hours", "grand_total"])

    for field_name in fields_to_apply:
        loc: WordLoc | None = getattr(det, field_name)
        value = values.get(field_name, "")
        if loc is None:
            unresolved.append(field_name)
            continue
        if loc.match_text:
            # Modern path: text-based find-and-replace.
            if not _replace_match_in_paragraphs(
                leaf_paragraphs, loc.match_text, value,
            ):
                unresolved.append(field_name)
        elif loc.table_index >= 0:
            # Legacy path for cached .billingapp.json files from older
            # versions — fall back to cell-based editing.
            try:
                cell = doc.tables[loc.table_index].rows[loc.row].cells[loc.col]
                _set_cell_text_legacy(cell, value, loc.paragraph_index)
            except IndexError:
                unresolved.append(field_name)
        else:
            unresolved.append(field_name)

    doc.save(str(path))
    return unresolved


def _adjust_formula_rows(ws, *, start_row: int, last_row: int, target_days: int) -> None:
    """Make sure exactly `target_days` data rows exist below the header.

    For shorter months, extra rows beyond target_days are blanked out so
    SUM() ranges that span the maximum month length still produce correct
    totals. For longer months we replicate the last existing data row,
    rewriting any formula references from the template-row index to the
    new row index.
    """
    import re as _re
    current_rows = last_row - start_row + 1
    needed_rows = target_days

    # Trim: blank out rows beyond target_days (within the existing range).
    if needed_rows < current_rows:
        for i in range(needed_rows, current_rows):
            row = start_row + i
            for col in range(1, ws.max_column + 1):
                ws.cell(row=row, column=col).value = None
        return

    # Extend: copy the template (last) row down to fill missing rows.
    if needed_rows > current_rows:
        ref_row = last_row
        max_col = ws.max_column
        for i in range(current_rows, needed_rows):
            new_row = start_row + i
            day_number = i + 1
            for col in range(1, max_col + 1):
                src = ws.cell(row=ref_row, column=col)
                dst = ws.cell(row=new_row, column=col)
                v = src.value
                if isinstance(v, str) and v.startswith("="):
                    # Bump every relative reference that points at ref_row
                    # forward to new_row (e.g. C36 -> C37).
                    new_v = _re.sub(
                        rf"(\$?[A-Z]+\$?){ref_row}\b",
                        rf"\g<1>{new_row}",
                        v,
                    )
                    dst.value = new_v
                else:
                    dst.value = v
                _copy_cell_style(src, dst)
            # Day-number column (col A): force a literal value so subsequent
            # rows referencing A_ via formula still chain correctly even if
            # the source row's A cell happened to be a literal.
            ws.cell(row=new_row, column=1).value = day_number


def _copy_cell_style(src, dst) -> None:
    from copy import copy as _copy
    if src.has_style:
        dst.font = _copy(src.font)
        dst.border = _copy(src.border)
        dst.fill = _copy(src.fill)
        dst.number_format = src.number_format
        dst.alignment = _copy(src.alignment)
        dst.protection = _copy(src.protection)


def _swap_year_in_formulas(ws, old_year: int, new_year: int) -> None:
    """Replace `old_year` with `new_year` everywhere it appears as a literal
    inside a formula string. Used when the user moves into a new calendar
    year (e.g. April 2026 -> January 2027).
    """
    import re as _re
    pattern = _re.compile(rf"(?<![0-9]){old_year}(?![0-9])")
    for row in ws.iter_rows():
        for cell in row:
            v = cell.value
            if isinstance(v, str) and v.startswith("=") and pattern.search(v):
                cell.value = pattern.sub(str(new_year), v)


def _set_cell_text_legacy(cell, text: str, paragraph_index: int = 0) -> None:
    if paragraph_index >= len(cell.paragraphs):
        paragraph_index = 0
    paragraph = cell.paragraphs[paragraph_index]
    if paragraph.runs:
        paragraph.runs[0].text = text
        for extra in paragraph.runs[1:]:
            extra.text = ""
    else:
        paragraph.add_run(text)


def _leaf_paragraph_elements(doc) -> list:
    """Return leaf <w:p> elements — paragraphs without nested <w:p>, matching
    what auto_detect walks over."""
    p_tag = qn("w:p")
    out = []
    for p in doc.element.body.iter(p_tag):
        has_nested = False
        for child in p.iter(p_tag):
            if child is not p:
                has_nested = True
                break
        if not has_nested:
            out.append(p)
    return out


def _replace_match_in_paragraphs(paragraphs, old: str, new: str) -> bool:
    """Replace every occurrence of `old` with `new` across run boundaries in
    any of the given paragraphs. Returns True if at least one replacement
    was made.
    """
    if not old:
        return False
    t_tag = qn("w:t")
    any_replaced = False
    for p in paragraphs:
        runs = list(p.iter(t_tag))
        if not runs:
            continue
        # Keep replacing until the joined text no longer contains `old`; this
        # lets us catch duplicated layouts that include the same value twice.
        while True:
            joined = "".join((t.text or "") for t in runs)
            idx = joined.find(old)
            if idx < 0:
                break
            end = idx + len(old)
            _splice_runs(runs, idx, end, new)
            any_replaced = True
    return any_replaced


def _splice_runs(runs, start: int, end: int, new: str) -> None:
    """Rewrite the given <w:t> run elements so that the joined character range
    [start, end) is replaced with `new`. Runs outside the range are
    untouched; runs fully inside are emptied; boundary runs keep their
    unaffected portions.
    """
    cursor = 0
    replaced = False
    for t in runs:
        txt = t.text or ""
        r_start = cursor
        r_end = cursor + len(txt)
        cursor = r_end
        if r_end <= start or r_start >= end:
            continue  # no overlap
        local_start = max(0, start - r_start)
        local_end = min(len(txt), end - r_start)
        prefix = txt[:local_start]
        suffix = txt[local_end:]
        if not replaced:
            t.text = prefix + new + suffix
            replaced = True
        else:
            # Middle / tail runs: keep only the portion outside the match.
            t.text = prefix + suffix


def _format_period(year: int, month: int) -> str:
    start, end = _format_period_endpoints(year, month)
    return f"{start} – {end}"


def _format_period_endpoints(year: int, month: int) -> tuple[str, str]:
    fmt = "%#m/%#d/%Y" if platform.system() == "Windows" else "%-m/%-d/%Y"
    start = date(year, month, 1).strftime(fmt)
    end = date(year, month, days_in_month(year, month)).strftime(fmt)
    return start, end


def _format_invoice_date(d: date) -> str:
    return d.strftime("%B %d, %Y").replace(" 0", " ")


def _format_number(value: float) -> str:
    if float(value).is_integer():
        return str(int(value))
    return f"{value:.2f}"


def _format_currency(value: float) -> str:
    if float(value).is_integer():
        return f"${int(value):,}"
    return f"${value:,.2f}"


def _col_to_index(letters: str) -> int:
    acc = 0
    for c in letters.upper():
        acc = acc * 26 + (ord(c) - ord("A") + 1)
    return acc
